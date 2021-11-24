import sqlite3
import pandas as pd
import numpy as np
import docx
from docx.enum import table

con = sqlite3.connect('test.db')


def run_time():
    sql_query = """SELECT*FROM testidprod
    where partner is NULL and state is NULL and bs=0 and (factor = 1 or factor = 2);"""
    dff = pd.read_sql(sql_query, con)
    dff = dff.drop(columns=['id', 'mir', 'raw', 'hash', 'meta', 'partner', 'state', 'bs', 'country'])
    dff = dff.reindex(columns=['factor', 'year', 'res'])

    year_add_dff = pd.DataFrame({'factor': [1, 1, 2, 2],    # создаю отдельный датафрейм и добавляю недостающие года с соответствующими значениямт
                                 'year': [2006, 2020, 2006, 2020],
                                 'res': [np.NaN, np.NaN, np.NaN, np.NaN]})

    df1 = dff.append(year_add_dff, ignore_index=True)  # присоединяю созданный датафрейм к основному

    factor_add_df = df1[1: 16] # создаю дополнительный датафрейм для фактора 6 с целью дальнайшего присоединения к основному датафрейму
    factor_add_df.loc[:, 'factor'] = 6
    factor_add_df.loc[:, 'year'] = range(2006, 2021)

    combined_df = [df1, factor_add_df]

    sums = pd.concat(combined_df, sort=False)
    sums = sums.groupby(['factor', 'year']).sum('res').rename(columns={'res': 'world'})

    factor2_values = sums['world'].loc[2]
    factor1_values = sums['world'].loc[1]
    factor6_values = factor2_values.div(factor1_values)

    values6 = factor6_values.array
    value_of_factor6 = []
    for value in values6:
        value_of_factor6.append(round(value, 3))

    sums.loc[6, 'world'] = value_of_factor6
    value_of_factor6_not_null = values6.dropna()
    finish_digit = round(value_of_factor6_not_null[-1], 3)
    start_digit = round(value_of_factor6_not_null[0], 3)

    s_year = factor6_values[factor6_values == value_of_factor6_not_null[0]].index[0]
    y_year = factor6_values[factor6_values == value_of_factor6_not_null[-1]].index[0]

    cagr = round((((finish_digit / start_digit) ** (1 / (y_year - s_year))) - 1) * 100, 2)
    if cagr > 0:
        direction = 'grew'
    else:
        direction = 'decreased'

    df = sums
    sums = sums.T

    sums.to_excel('report.xlsx')

    df = df.loc[6]
    df.loc[:, 'factor'] = 6
    df.loc[:, 'year'] = range(2006, 2021)
    df = df.rename(columns={'world': 'world value'})
    df = df.reindex(columns=['factor', 'year', 'world value'])

    doc = docx.Document()
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j].title()
        t.cell(0, j).paragraphs[0].runs[0].font.bold = True

    u = t.cell(1, 0)

    u.alignment = table.WD_TABLE_ALIGNMENT.CENTER
    u.vertical_alignment = table.WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if j == 0:
                if i != 1:
                    u.merge(t.cell(i + 1, j))
            if j < 2:
                t.cell(i + 1, j).text = str(int(df.values[i, j]))
            else:
                t.cell(i + 1, j).text = str(round(df.values[i, j], 2))

    doc.add_paragraph(f'Factor 6 {direction}  by avg {cagr}% every year from {s_year} to {y_year}')
    doc.save('report.docx')


if __name__ == '__main__':
    run_time()
